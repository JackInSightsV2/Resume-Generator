#!/usr/bin/env python3
"""
Convert a Markdown resume into a formatted Word document.

This version:
  - Automatically converts Markdown headings (e.g. "#", "##", "###", etc.) into
    Word heading styles.
  - Converts inline markdown for bold (**text**) and italics (*text*) into formatted text.
  - Converts markdown hyperlinks [link text](url) into clickable links.
  - Sets the page size and margins.
  - Renders horizontal rules (lines containing only '---') as horizontal lines.
  - In the Certifications table:
      • The header row gets a pastel green background with dark green text.
      • All cells in the data rows (all rows except the title row) get a pastel yellow background.
  - Uses the Aptos font throughout.
  
Usage:
    python convert_resume.py path/to/input.md output.docx [--path /resumes]
    (If --path is not provided, the output file is saved in the exported_resumes folder.)
"""

import re
import sys
import argparse
import os

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.opc.constants import RELATIONSHIP_TYPE

# -------------------------------
# Helper functions
# -------------------------------

def set_cell_background(cell, color_hex):
    """
    Set the background shading of a cell.
    color_hex: a string like "FFFACD" (no '#' character).
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def add_hyperlink(paragraph, text, url):
    """
    Add a hyperlink to the given paragraph.
    The link will be underlined and use the Aptos font.
    """
    part = paragraph.part
    r_id = part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    
    # Create the hyperlink element.
    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)
    
    # Create a run for the link text.
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    # Underline the text.
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)
    # Set the font to Aptos.
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), "Aptos")
    rFonts.set(qn('w:hAnsi'), "Aptos")
    rPr.append(rFonts)
    
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)
    return hyperlink

def insert_markdown_text(paragraph, markdown_text):
    """
    Parse the markdown_text and add runs to the given paragraph.
    Handles:
      - Bold: **text**
      - Italic: *text*
      - Hyperlinks: [link text](url)
    All text uses the Aptos font.
    """
    # Regex to match either **bold**, *italic*, or [link text](url)
    pattern = re.compile(r"(\*\*(?P<bold>.+?)\*\*|\*(?P<italic>.+?)\*|\[(?P<link_text>[^\]]+)\]\((?P<link_url>[^)]+)\))")
    pos = 0
    for match in pattern.finditer(markdown_text):
        # Add any text before the match.
        if match.start() > pos:
            run = paragraph.add_run(markdown_text[pos:match.start()])
            run.font.name = "Aptos"
        if match.group('bold'):
            run = paragraph.add_run(match.group('bold'))
            run.bold = True
            run.font.name = "Aptos"
        elif match.group('italic'):
            run = paragraph.add_run(match.group('italic'))
            run.italic = True
            run.font.name = "Aptos"
        elif match.group('link_text'):
            add_hyperlink(paragraph, match.group('link_text'), match.group('link_url'))
        pos = match.end()
    if pos < len(markdown_text):
        run = paragraph.add_run(markdown_text[pos:])
        run.font.name = "Aptos"

def insert_formatted_text(paragraph, text):
    """
    Convenience function to insert markdown-formatted text into a paragraph.
    """
    insert_markdown_text(paragraph, text)

def add_paragraph_with_formatting(doc, text, style=None, alignment=None):
    """
    Add a paragraph to the document with the given text (processing inline markdown).
    """
    p = doc.add_paragraph(style=style)
    if alignment is not None:
        p.alignment = alignment
    insert_formatted_text(p, text)
    return p

def add_horizontal_line(doc):
    """
    Insert a horizontal line by adding a paragraph with a bottom border.
    """
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p_pr = p._p.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "auto")
    p_bdr.append(bottom)
    p_pr.append(p_bdr)

def try_process_heading(doc, line, header_block=False):
    """
    If the line is a Markdown heading (one or more '#' followed by a space), add a
    corresponding Word heading paragraph.
    
    header_block: if True, center the paragraph (used for the top block).
    
    Returns a tuple (is_heading, level, text) where:
      - is_heading: True if the line was processed as a heading.
      - level: the heading level (1 to 6) if processed.
      - text: the heading text.
    """
    m = re.match(r'^(#{1,6})\s+(.*)$', line)
    if m:
        level = len(m.group(1))
        text = m.group(2).strip()
        if level == 1:
            style = 'Heading1'
        elif level == 2:
            style = 'Heading2'
        elif level == 3:
            style = 'Heading3'
        else:
            style = 'Normal'
        p = doc.add_paragraph(style=style)
        if header_block:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        insert_formatted_text(p, text)
        return True, level, text
    return False, None, None

def add_job_block_to_doc(doc, block_lines):
    """
    Add a block of lines (for an employment/job entry) to the document.
    Sets "keep with next" on all paragraphs except the last to help keep them together on one page.
    """
    paragraphs = []
    for line in block_lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith("- "):
            p = doc.add_paragraph(style='List Bullet')
            insert_formatted_text(p, line[2:].strip())
        else:
            p = doc.add_paragraph()
            insert_formatted_text(p, line)
        paragraphs.append(p)
    for p in paragraphs[:-1]:
        p.paragraph_format.keep_with_next = True

def add_table_to_doc(doc, table_lines):
    """
    Convert a Markdown table (list of lines) into a Word table.
      - The header row gets a pastel green background with dark green text.
      - All cells in the data rows (rows below the header) get a pastel yellow background with black text.
    """
    if len(table_lines) < 2:
        return

    # Process header row.
    header_line = table_lines[0]
    header_cells_text = [cell.strip() for cell in header_line.strip("|").split("|")]
    # Data rows: skip the separator (second line) and process the rest.
    data_lines = table_lines[2:]
    
    table = doc.add_table(rows=len(data_lines) + 1, cols=len(header_cells_text))
    table.style = 'Table Grid'

    # Header row formatting.
    hdr_cells = table.rows[0].cells
    for j, text in enumerate(header_cells_text):
        p = hdr_cells[j].paragraphs[0]
        p.text = ""
        insert_formatted_text(p, text)
        set_cell_background(hdr_cells[j], "C6EFCE")  # Pastel green.
        for paragraph in hdr_cells[j].paragraphs:
            for run in paragraph.runs:
                run.font.color.rgb = RGBColor(0x00, 0x61, 0x00)  # Dark green.
                run.font.name = "Aptos"

    # Data rows: Apply pastel yellow background to every cell.
    for i, line in enumerate(data_lines):
        cells = [cell.strip() for cell in line.strip("|").split("|")]
        row_cells = table.rows[i+1].cells
        for j, text in enumerate(cells):
            p = row_cells[j].paragraphs[0]
            p.text = ""
            insert_formatted_text(p, text)
            set_cell_background(row_cells[j], "FFFACD")  # Pastel yellow.
            for paragraph in row_cells[j].paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)  # Black text.
                    run.font.name = "Aptos"

# -------------------------------
# Main conversion function
# -------------------------------

def convert_md_to_docx(md_file, docx_file, output_path=None):
    # Determine the output file location.
    if output_path:
        # If the path is not absolute, make it absolute relative to the current working directory.
        if not os.path.isabs(output_path):
            output_path = os.path.abspath(output_path)
        # Create the folder if it doesn't exist.
        if not os.path.exists(output_path):
            os.makedirs(output_path)
        # Place the output file in the specified folder.
        docx_file = os.path.join(output_path, os.path.basename(docx_file))
    else:
        # If no output path is provided and no directory is specified in the docx_file, 
        # default to the exported_resumes folder.
        if not os.path.dirname(docx_file):
            exported_folder = os.path.join(os.path.dirname(__file__), "exported_resumes")
            if not os.path.exists(exported_folder):
                os.makedirs(exported_folder)
            docx_file = os.path.join(exported_folder, docx_file)
        else:
            docx_file = os.path.abspath(docx_file)

    with open(md_file, 'r', encoding='utf-8') as f:
        md_text = f.read()

    # Remove code fence markers if present at the top and bottom.
    lines = md_text.splitlines()
    if lines and lines[0].strip() == "```markdown":
        lines = lines[1:]
    if lines and lines[-1].strip() == "```":
        lines = lines[:-1]
    md_text = "\n".join(lines)

    doc = Document()

    # Set page dimensions and margins.
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    # Set margins to 0.5 inches on all sides.
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.5)
    section.bottom_margin = Inches(0.5)

    # Set default style: single spacing, no extra spacing, Aptos font.
    style = doc.styles['Normal']
    style.font.name = "Aptos"
    style.font.size = Pt(11)
    para_format = style.paragraph_format
    para_format.line_spacing = 1
    para_format.space_before = Pt(0)
    para_format.space_after = Pt(0)

    # -------------------------------
    # Read header text from header.txt file.
    # -------------------------------
    header_path = os.path.join(os.path.dirname(__file__), "settings", "header.txt")
    try:
        with open(header_path, 'r', encoding='utf-8') as hf:
            header_text = hf.read().strip()
    except Exception as e:
        sys.exit(f"Error reading header file at {header_path}: {e}")

    # Add document header (applied on every page).
    header = section.header
    if header.paragraphs:
        hp = header.paragraphs[0]
        hp.clear()
    else:
        hp = header.add_paragraph()
    insert_markdown_text(hp, header_text)
    hp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Split Markdown into lines.
    lines = md_text.splitlines()
    i = 0

    # Process the top header block (everything until the first horizontal rule).
    header_block_lines = []
    while i < len(lines) and lines[i].strip() != "---":
        header_block_lines.append(lines[i])
        i += 1

    # Process header block lines.
    for line in header_block_lines:
        stripped = line.strip()
        if not stripped:
            continue
        is_heading, level, text = try_process_heading(doc, stripped, header_block=True)
        if not is_heading:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            insert_formatted_text(p, stripped)

    # Process any horizontal rule(s) in the header block.
    while i < len(lines) and lines[i].strip() == "---":
        add_horizontal_line(doc)
        i += 1

    # Process the remaining content.
    in_employment = False  # Flag to indicate we're inside the EMPLOYMENT HISTORY block.
    current_job_block = []  # Collect lines belonging to a single job block.
    table_lines = []        # Collect lines of a table.

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Render horizontal rules.
        if stripped == "---":
            if in_employment and current_job_block:
                add_job_block_to_doc(doc, current_job_block)
                current_job_block = []
            add_horizontal_line(doc)
            i += 1
            continue

        # Check if the line is a Markdown heading.
        is_heading, level, heading_text = try_process_heading(doc, stripped, header_block=False)
        if is_heading:
            # If we were inside an employment block, flush it.
            if in_employment and current_job_block:
                add_job_block_to_doc(doc, current_job_block)
                current_job_block = []
            # If this is a level‑2 heading and its text is "EMPLOYMENT HISTORY", set the flag.
            if level == 2 and heading_text.upper() == "EMPLOYMENT HISTORY":
                in_employment = True
            else:
                in_employment = False
            i += 1
            continue

        # Check for table lines (assume lines starting with "|" are part of a table).
        if stripped.startswith("|"):
            table_lines.append(line)
            # If the next line isn’t part of the table, flush the table.
            if i + 1 < len(lines) and not lines[i+1].strip().startswith("|"):
                add_table_to_doc(doc, table_lines)
                table_lines = []
            elif i + 1 == len(lines):
                add_table_to_doc(doc, table_lines)
                table_lines = []
            i += 1
            continue

        # Process normal text lines.
        if in_employment:
            current_job_block.append(line)
        else:
            if stripped.startswith("- "):
                p = doc.add_paragraph(style='List Bullet')
                insert_formatted_text(p, stripped[2:].strip())
            else:
                p = doc.add_paragraph()
                insert_formatted_text(p, line)
        i += 1

    # Flush any remaining job block or table.
    if current_job_block:
        add_job_block_to_doc(doc, current_job_block)
    if table_lines:
        add_table_to_doc(doc, table_lines)

    doc.save(docx_file)
    print(f"Saved formatted resume to {docx_file}")

# -------------------------------
# Command-line interface
# -------------------------------

def main():
    parser = argparse.ArgumentParser(
        description='Convert a Markdown resume into a formatted Word document.'
    )
    parser.add_argument('input', help='Input Markdown file')
    parser.add_argument('output', help='Output DOCX file name (without a folder path if using --path)')
    parser.add_argument('--path', type=str, default=None,
                        help='Output folder path where the DOCX file will be saved. '
                             'If the folder does not exist, it will be created. '
                             'If not provided, the file is saved in the exported_resumes folder.')
    args = parser.parse_args()
    convert_md_to_docx(args.input, args.output, args.path)

if __name__ == '__main__':
    main()